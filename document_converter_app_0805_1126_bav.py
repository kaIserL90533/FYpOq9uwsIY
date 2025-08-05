# 代码生成时间: 2025-08-05 11:26:56
import os
from django.conf import settings
from django.http import JsonResponse
from django.views import View
from django.core.files.storage import default_storage
from docx import Document
from openpyxl import load_workbook
from pdf2docx import Converter


class DocumentConverterView(View):
    """
    A Django view to convert documents between formats.
    Supported formats: PDF to DOCX, DOCX to XLSX, and vice versa.
    """
    def post(self, request):
        """
        Handle the POST request to convert documents.
        Expects a JSON payload containing file information and the desired conversion.
        """
        try:
            file_id = request.POST.get('file_id')
            conversion_type = request.POST.get('conversion_type')

            # Check for valid file_id and conversion_type
            if not file_id or not conversion_type:
                return JsonResponse({'error': 'Missing file_id or conversion_type'}, status=400)

            file_path = os.path.join(settings.MEDIA_ROOT, 'uploads', file_id)

            if not os.path.exists(file_path):
                return JsonResponse({'error': 'File not found'}, status=404)

            # Perform the conversion based on the conversion_type
            if conversion_type == 'pdf_to_docx':
                return convert_pdf_to_docx(file_path)
            elif conversion_type == 'docx_to_xlsx':
                return convert_docx_to_xlsx(file_path)
            elif conversion_type == 'xlsx_to_docx':
                return convert_xlsx_to_docx(file_path)
            else:
                return JsonResponse({'error': 'Unsupported conversion type'}, status=400)

        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)


def convert_pdf_to_docx(pdf_path):
    """Convert a PDF file to a DOCX file."""
    try:
        converter = Converter(pdf_path)
        output_path = os.path.splitext(pdf_path)[0] + '.docx'
        converter.convert(output_path, start=0, end=None)
        converter.close()
        return JsonResponse({'message': 'Conversion successful', 'file_path': output_path})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def convert_docx_to_xlsx(docx_path):
    """Convert a DOCX file to an XLSX file."""
    try:
        document = Document(docx_path)
        wb = load_workbook()
        ws = wb.active
        for para in document.paragraphs:
            ws.append([para.text])
        output_path = os.path.splitext(docx_path)[0] + '.xlsx'
        wb.save(output_path)
        return JsonResponse({'message': 'Conversion successful', 'file_path': output_path})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def convert_xlsx_to_docx(xlsx_path):
    """Convert an XLSX file to a DOCX file."""
    try:
        wb = load_workbook(xlsx_path)
        ws = wb.active
        document = Document()
        for row in ws.iter_rows(min_row=1, max_col=1, values_only=True):
            document.add_paragraph("
".join(str(cell) for cell in row))
        output_path = os.path.splitext(xlsx_path)[0] + '.docx'
        document.save(output_path)
        return JsonResponse({'message': 'Conversion successful', 'file_path': output_path})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
